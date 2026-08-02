import sys
if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

"""
docx_processor - Word _

_
  1. _ Word _
  2. _
  3. _ AI _ 300 _Doc summary_
  4. _ extra_description _

_
  - python-docx (Word _)
  - zhipuai SDK (_)
  - upload_test (_)

_
  from docx_processor import process_docx
  process_docx("report.docx", uploader="_")
"""

import os
import base64
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from dotenv import load_dotenv
from zhipuai import ZhipuAI

# [CN]
load_dotenv()
API_KEY = os.getenv("ZHIPU_API_KEY")
SUMMARY_MODEL = os.getenv("ZHIPU_SUMMARY_MODEL", "GLM-5V-Turbo")

if not API_KEY:
    raise RuntimeError("ZHIPU_API_KEY not found in .env")

client = ZhipuAI(api_key=API_KEY)

# [CN]
BASE_DIR = Path(__file__).parent
UPLOADS_DIR = BASE_DIR / "uploads"
TEMP_DIR = BASE_DIR / "temp_docx"


def extract_images_from_docx(docx_path: str | Path) -> list[Path]:
    """_ Word _ uploads/ _

    _IMAGE _
    _

    Args:
        docx_path: Word _ (.docx)_

    Returns:
        list[Path]: _
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Document not found: {docx_path}")

    doc = Document(docx_path)
    image_paths = []

    # [CN]
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob

            # [CN]
            target_name = Path(rel.target_part.partname).name
            suffix = Path(target_name).suffix or ".png"

            # [CN] uploads/[CN]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            dest_name = f"docx_{timestamp}_{len(image_paths)+1}{suffix}"
            dest_path = UPLOADS_DIR / dest_name

            with open(dest_path, "wb") as f:
                f.write(image_data)

            image_paths.append(dest_path)
            print(f"   Extracted: {dest_name}")

    # Release reference for garbage collection (Windows lock fix)
    del doc
    return image_paths


def extract_text_from_docx(docx_path: str | Path) -> str:
    """_ Word _

    _

    Args:
        docx_path: Word _ (.docx)_

    Returns:
        str: _
    """
    docx_path = Path(docx_path)
    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def generate_docx_summary(docx_path: str | Path, max_chars: int = 300) -> str:
    """_ AI _ Word _

    _ AI _
    _

    Args:
        docx_path: Word _ (.docx)_
        max_chars: _ 300_

    Returns:
        str: Doc summary_
    """
    docx_path = Path(docx_path)
    full_text = extract_text_from_docx(docx_path)

    if not full_text.strip():
        return "(no text content)"

    # [CN] token[CN]
    text_for_summary = full_text[:3000] if len(full_text) > 3000 else full_text

    try:
        response = client.chat.completions.create(
            model=SUMMARY_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": f"你是一个新闻Doc summary助手。请根据文档内容生成一段{max_chars}字以内的中文摘要，"
                           f"_"
                           f"_"
                },
                {
                    "role": "user",
                    "content": text_for_summary
                }
            ]
        )
        summary = response.choices[0].message.content
        return summary
    except Exception as e:
        print(f"   Warning: AI summary generation failed ({e})_300_")
        return full_text[:max_chars]


def process_docx(docx_path: str, uploader: str = "test_user", source_name: str = "") -> dict:
    """_ Word _ + _ + _

    _
      1. _ Word_
      2. _
      3. AI _ 300 _Doc summary
      4. _ process_image()_ extra_description
         _

    Args:
        docx_path: Word _ (.docx)_
        uploader: _

    Returns:
        dict: _ summary, image_count, image_results _
    """
    from upload_test import process_image
    from supabase_client import upload_to_storage

    docx_path = Path(docx_path)
    docx_name = docx_path.stem

    print(f"\n{'='*60}")
    print(f"Processing DOCX: {docx_name}")
    print(f"{'='*60}")

    # 0. [CN] DOCX [CN] Supabase Storage
    print("\n[Step 0] Uploading original DOCX to Storage...")
    try:
        source_url = upload_to_storage("uploads", docx_path)
        print(f"   Source URL: {source_url}")
    except Exception as e:
        print(f"   Warning: DOCX upload failed ({e}), skipping source link")
        source_url = ""

    # 1. [CN]
    print("\n[Step 1] Extracting images...")
    image_paths = extract_images_from_docx(docx_path)
    print(f"   Found {len(image_paths)} images")

    # 2. [CN]
    print("\n[Step 2] Extracting text...")
    full_text = extract_text_from_docx(docx_path)
    text_preview = full_text[:100] + "..." if len(full_text) > 100 else full_text
    print(f"   Text preview: {text_preview}")

    # 3. [CN]Doc summary
    print("\n[Step 3] Generating summary...")
    summary = generate_docx_summary(docx_path, max_chars=300)
    print(f"   Summary: {summary}")

    # 4. [CN]Doc summary + [CN]
    source_label = source_name if source_name else f"{docx_name}.docx"
    print(f"\n[Step 4] Indexing {len(image_paths)} images with summary...")
    image_results = []
    for img_path in image_paths:
        try:
            result = process_image(
                str(img_path),
                uploader=uploader,
                extra_description=f"[Doc summary] {summary}",
                source_file=source_label,
                source_url=source_url,
            )
            image_results.append(result)
        except Exception as e:
            print(f"   Failed: {img_path.name} - {e}")

    print(f"\n{'='*60}")
    print(f"DOCX processing complete: {len(image_results)}/{len(image_paths)} images indexed")
    print(f"{'='*60}")

    return {
        "docx_name": docx_name,
        "summary": summary,
        "image_count": len(image_paths),
        "image_results": image_results
    }


# ============================================================
# [CN]
# ============================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_processor.py <docx_file_path> [uploader_name]")
        print("Example: python docx_processor.py report.docx _")
        sys.exit(1)

    docx_file = sys.argv[1]
    uploader_name = sys.argv[2] if len(sys.argv) > 2 else "test_user"

    result = process_docx(docx_file, uploader=uploader_name)
    print(f"\nFinal result: {result['image_count']} images indexed")
    print(f"Summary: {result['summary']}")